"""Deep comparison of original template vs generated resume - every detail."""
from docx import Document
from docx.oxml.ns import qn
import sys


def get_para_details(para):
    """Extract every formatting detail from a paragraph."""
    pPr = para._element.find(qn('w:pPr'))
    
    # Check for borders, spacing, indentation
    details = {
        "style": para.style.name if para.style else None,
        "alignment": str(para.alignment),
        "text": para.text,
        "num_runs": len(para.runs),
        "run_texts": [r.text for r in para.runs],
    }
    
    # Paragraph-level XML attributes
    if pPr is not None:
        pBdr = pPr.find(qn('w:pBdr'))
        details["has_border"] = pBdr is not None
        if pBdr is not None:
            details["border_xml"] = pBdr.xml if hasattr(pBdr, 'xml') else str(pBdr)
        
        ind = pPr.find(qn('w:ind'))
        details["indent"] = ind.attrib if ind is not None else None
        
        spacing = pPr.find(qn('w:spacing'))
        details["spacing"] = spacing.attrib if spacing is not None else None
        
        tabs = pPr.find(qn('w:tabs'))
        if tabs is not None:
            details["tabs"] = [t.attrib for t in tabs.findall(qn('w:tab'))]
        else:
            details["tabs"] = None
    
    # Run-level details
    run_details = []
    for r in para.runs:
        rd = {
            "text": r.text,
            "bold": r.bold,
            "italic": r.italic,
            "underline": r.underline,
            "font_name": r.font.name,
            "font_size": str(r.font.size) if r.font.size else None,
            "font_color": str(r.font.color.rgb) if r.font.color and r.font.color.rgb else None,
        }
        run_details.append(rd)
    details["runs"] = run_details
    
    return details


def compare_docs(orig_path, gen_path):
    orig = Document(orig_path)
    gen = Document(gen_path)
    
    print(f"Original: {len(orig.paragraphs)} paragraphs")
    print(f"Generated: {len(gen.paragraphs)} paragraphs")
    print()
    
    if len(orig.paragraphs) != len(gen.paragraphs):
        print("!!! PARAGRAPH COUNT MISMATCH !!!")
        print()
    
    issues = []
    
    for i in range(max(len(orig.paragraphs), len(gen.paragraphs))):
        if i >= len(orig.paragraphs):
            issues.append(f"\n[{i}] EXTRA paragraph in generated: \"{gen.paragraphs[i].text[:80]}\"")
            continue
        if i >= len(gen.paragraphs):
            issues.append(f"\n[{i}] MISSING paragraph in generated (was: \"{orig.paragraphs[i].text[:80]}\")")
            continue
        
        o = get_para_details(orig.paragraphs[i])
        g = get_para_details(gen.paragraphs[i])
        
        para_issues = []
        
        if o["style"] != g["style"]:
            para_issues.append(f"  STYLE: '{o['style']}' -> '{g['style']}'")
        if o["alignment"] != g["alignment"]:
            para_issues.append(f"  ALIGNMENT: {o['alignment']} -> {g['alignment']}")
        if o.get("has_border") != g.get("has_border"):
            para_issues.append(f"  BORDER: {o.get('has_border')} -> {g.get('has_border')}")
        if o.get("indent") != g.get("indent"):
            para_issues.append(f"  INDENT: {o.get('indent')} -> {g.get('indent')}")
        if o.get("spacing") != g.get("spacing"):
            para_issues.append(f"  SPACING: {o.get('spacing')} -> {g.get('spacing')}")
        if o.get("tabs") != g.get("tabs"):
            para_issues.append(f"  TABS: {o.get('tabs')} -> {g.get('tabs')}")
        if o["num_runs"] != g["num_runs"]:
            para_issues.append(f"  RUN COUNT: {o['num_runs']} -> {g['num_runs']}")
            para_issues.append(f"    orig runs: {o['run_texts']}")
            para_issues.append(f"    gen  runs: {g['run_texts']}")
        
        # Compare run formatting
        for j in range(min(len(o["runs"]), len(g["runs"]))):
            or_ = o["runs"][j]
            gr_ = g["runs"][j]
            for key in ["bold", "italic", "underline", "font_name", "font_size", "font_color"]:
                if or_[key] != gr_[key]:
                    para_issues.append(f"  RUN[{j}].{key}: {or_[key]} -> {gr_[key]}")
        
        if para_issues:
            header = f"\n[{i}] ({o['style']}) \"{o['text'][:60]}\""
            issues.append(header)
            for p in para_issues:
                issues.append(p)
    
    # Check tables
    print(f"Original tables: {len(orig.tables)}")
    print(f"Generated tables: {len(gen.tables)}")
    
    # Check sections (page margins, etc)
    for si, (os, gs) in enumerate(zip(orig.sections, gen.sections)):
        sec_issues = []
        for attr in ["left_margin", "right_margin", "top_margin", "bottom_margin", 
                      "page_width", "page_height", "orientation"]:
            ov = getattr(os, attr, None)
            gv = getattr(gs, attr, None)
            if ov != gv:
                sec_issues.append(f"  Section[{si}].{attr}: {ov} -> {gv}")
        if sec_issues:
            issues.append(f"\nSECTION {si} DIFFERENCES:")
            issues.extend(sec_issues)
    
    if issues:
        print(f"\n{'='*60}")
        print(f"FOUND {len([x for x in issues if x.startswith(chr(10))])} ISSUES:")
        print(f"{'='*60}")
        for iss in issues:
            print(iss)
    else:
        print("\nNO ISSUES FOUND - Perfect match!")


if __name__ == "__main__":
    compare_docs(
        r"template\resume_template.docx",
        r"output\Rehan_Malik_Resume_data_engineer_amazon.docx"
    )
