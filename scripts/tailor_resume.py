"""
Tailor a resume to a job description using GitHub Models API,
preserving the exact DOCX structure, formatting, and styles.
"""

import os
import sys
import json
import copy
import argparse
import re
from pathlib import Path
from copy import deepcopy

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree


GITHUB_MODELS_URL = "https://models.inference.ai.azure.com/chat/completions"
MODEL_NAME = "gpt-4o"


def read_jd(jd_path: str) -> str:
    ext = Path(jd_path).suffix.lower()
    if ext == ".docx":
        doc = Document(jd_path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".pdf":
        try:
            import fitz
            with fitz.open(jd_path) as pdf:
                return "\n".join(page.get_text() for page in pdf)
        except ImportError:
            sys.exit("PyMuPDF (fitz) required for PDF JDs. Install with: pip install PyMuPDF")
    else:
        with open(jd_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


def build_experience_blocks(doc: Document) -> list[dict]:
    """
    Parse the resume into experience blocks: each block has a company header
    and a list of bullet-point paragraph indices (only non-empty bullets).
    """
    blocks = []
    current_block = None
    in_experience = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""

        if style == "Heading 1" and text.lower() == "experience":
            in_experience = True
            continue

        if style == "Heading 1" and ("education" in text.lower() or "certification" in text.lower()):
            in_experience = False
            if current_block:
                blocks.append(current_block)
                current_block = None
            continue

        if not in_experience:
            continue

        is_date_company = (
            style in ("Normal", "Body Text")
            and "\t" in text
            and not text.startswith("-")
            and text  # non-empty
        )
        if is_date_company:
            if current_block:
                blocks.append(current_block)
            current_block = {
                "date_company_idx": i,
                "title_idx": None,
                "bullet_indices": [],
                "company_text": text,
            }
            continue

        if style == "Heading 1" and current_block and current_block["title_idx"] is None:
            current_block["title_idx"] = i
            current_block["title_text"] = text
            continue

        if style == "List Paragraph" and current_block and text:
            current_block["bullet_indices"].append(i)
            continue

    if current_block:
        blocks.append(current_block)

    return blocks


def build_skills_entries(doc: Document) -> list[dict]:
    """
    Extract skill category entries from the skills section.
    Each entry maps a paragraph index to its category label and skills text.
    Continuation lines (no colon, not a heading) are appended to the previous entry.
    """
    in_skills = False
    entries = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""

        if style == "Heading 1" and "certification" in text.lower():
            in_skills = True
            continue

        if not in_skills:
            continue

        if not text:
            continue

        # Skip heading/label-only paragraphs like "skills"
        if style == "Normal" and ":" not in text and len(text) < 20:
            continue

        bold_part = ""
        normal_part = ""
        found_colon = False

        for run in para.runs:
            if not found_colon:
                if ":" in run.text:
                    colon_idx = run.text.index(":")
                    bold_part += run.text[:colon_idx + 1]
                    normal_part += run.text[colon_idx + 1:]
                    found_colon = True
                else:
                    bold_part += run.text
            else:
                normal_part += run.text

        if bold_part.strip() and normal_part.strip():
            entries.append({
                "para_idx": i,
                "continuation_indices": [],
                "category": bold_part.strip(),
                "skills": normal_part.strip(),
                "full_text": text,
            })
        elif entries and not found_colon and text:
            # Continuation line - append skills to previous entry
            entries[-1]["skills"] += " " + text
            entries[-1]["continuation_indices"].append(i)

    return entries


def call_github_model(token: str, system_prompt: str, user_prompt: str) -> str:
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.4,
        "max_tokens": 4000,
    }
    resp = requests.post(GITHUB_MODELS_URL, headers=headers, json=payload, timeout=120)
    if resp.status_code != 200:
        print(f"API Error {resp.status_code}: {resp.text[:500]}")
        resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def tailor_bullets(token: str, jd_text: str, company: str, title: str,
                   original_bullets: list[str]) -> list[str]:
    system_prompt = (
        "You are an expert resume writer. You tailor resume bullet points to match "
        "a target job description while keeping them truthful and grounded in the "
        "original experience. Preserve technical depth and quantified metrics. "
        "Do NOT invent new experiences or metrics that aren't in the original. "
        "Do NOT add any prefix numbering or dashes. "
        "Do NOT use any markdown formatting (no **, no *, no #). "
        "Return ONLY plain text bullet points, one per line, "
        "with no extra commentary. Keep the EXACT same number of bullets."
    )
    user_prompt = (
        f"TARGET JOB DESCRIPTION:\n{jd_text}\n\n"
        f"CURRENT ROLE: {title} at {company}\n\n"
        f"ORIGINAL BULLETS ({len(original_bullets)} total - return EXACTLY {len(original_bullets)} lines):\n"
        + "\n".join(f"- {b}" for b in original_bullets)
        + f"\n\nRewrite each bullet to better align with the target JD. "
        f"Return EXACTLY {len(original_bullets)} lines of plain text, no dashes or numbers."
    )

    result = call_github_model(token, system_prompt, user_prompt)
    lines = [ln.strip().lstrip("-•0123456789. ") for ln in result.strip().split("\n") if ln.strip()]
    lines = [re.sub(r'\*\*([^*]+)\*\*', r'\1', ln) for ln in lines]

    if len(lines) != len(original_bullets):
        if len(lines) > len(original_bullets):
            lines = lines[:len(original_bullets)]
        else:
            lines.extend(original_bullets[len(lines):])

    return lines


def tailor_skills_entries(token: str, jd_text: str,
                          entries: list[dict]) -> list[dict]:
    """
    Ask the LLM to reorder skills within each category, and reorder
    the categories themselves. Returns list of {category, skills} dicts.
    """
    system_prompt = (
        "You are an expert resume writer. You reorder skills within each category "
        "of a resume to better match a target job description. "
        "Do NOT invent skills the candidate doesn't have. "
        "You may reorder skills WITHIN each category to put the most relevant ones first. "
        "You may reorder the CATEGORIES themselves to put the most relevant first. "
        "Keep the EXACT same category names and the EXACT same number of categories. "
        "Do NOT use any markdown formatting (no **, no *, no #, no backticks). "
        "Return ONLY plain text, one category per line, in the format: "
        "CategoryName: skill1, skill2, skill3, ..."
    )

    cat_lines = []
    for e in entries:
        cat_lines.append(f"{e['category']} {e['skills']}")

    user_prompt = (
        f"TARGET JOB DESCRIPTION:\n{jd_text}\n\n"
        f"ORIGINAL SKILLS ({len(entries)} categories - return EXACTLY {len(entries)} lines):\n"
        + "\n".join(cat_lines)
        + f"\n\nReorder skills within each category and reorder the categories. "
        f"Return EXACTLY {len(entries)} lines. "
        f"Format: CategoryName: skill1, skill2, ..."
    )

    result = call_github_model(token, system_prompt, user_prompt)
    new_lines = [ln.strip() for ln in result.strip().split("\n") if ln.strip()]
    new_lines = [re.sub(r'\*\*([^*]+)\*\*', r'\1', ln) for ln in new_lines]
    new_lines = [ln.strip('`') for ln in new_lines]

    new_entries = []
    for ln in new_lines:
        if ":" in ln:
            colon_idx = ln.index(":")
            cat = ln[:colon_idx + 1].strip()
            skills = ln[colon_idx + 1:].strip()
            new_entries.append({"category": cat, "skills": skills})

    if len(new_entries) != len(entries):
        print(f"  WARNING: LLM returned {len(new_entries)} categories, expected {len(entries)}. Using original.")
        return entries

    return new_entries


def replace_bullet_text(para, new_text: str):
    """
    Replace bullet text while preserving the multi-run structure.
    Distributes new text across runs in the same word-per-run pattern
    as the original to maintain identical rendering.
    """
    if not para.runs:
        para.text = new_text
        return

    if len(para.runs) == 1:
        para.runs[0].text = new_text
        return

    original_runs = para.runs
    new_words = new_text.split(" ")

    # The original pattern alternates: word, " ", word, " ", ... lastwords
    # We rebuild: put one word per pair of runs (word-run + space-run),
    # with remaining text in the last "word" run.
    word_run_indices = []
    space_run_indices = []
    for ri, run in enumerate(original_runs):
        if run.text.strip():
            word_run_indices.append(ri)
        elif run.text == " ":
            space_run_indices.append(ri)

    if not word_run_indices:
        original_runs[0].text = new_text
        for run in original_runs[1:]:
            run.text = ""
        return

    # Distribute words across word-runs, put overflow in the last word-run
    for ri, run in enumerate(original_runs):
        run.text = ""

    if len(new_words) <= len(word_run_indices):
        for wi, word in enumerate(new_words):
            original_runs[word_run_indices[wi]].text = word
        for si in range(len(new_words) - 1):
            if si < len(space_run_indices):
                original_runs[space_run_indices[si]].text = " "
    else:
        for wi in range(len(word_run_indices) - 1):
            original_runs[word_run_indices[wi]].text = new_words[wi]
        overflow = " ".join(new_words[len(word_run_indices) - 1:])
        original_runs[word_run_indices[-1]].text = overflow
        for si in range(len(word_run_indices) - 1):
            if si < len(space_run_indices):
                original_runs[space_run_indices[si]].text = " "


def replace_skills_text(para, new_category: str, new_skills: str):
    """
    Replace skills paragraph text while preserving the bold category name
    and normal-weight skills list pattern.

    Uses the actual bold property of each run to determine where to place
    category text (in bold runs) vs skills text (in non-bold runs).
    """
    if not para.runs:
        para.text = f"{new_category} {new_skills}"
        return

    # Classify runs by actual bold formatting
    first_bold_idx = None
    first_nonbold_idx = None
    for ri, run in enumerate(para.runs):
        if run.bold and first_bold_idx is None:
            first_bold_idx = ri
        if not run.bold and first_nonbold_idx is None:
            first_nonbold_idx = ri

    # Clear all run text
    for run in para.runs:
        run.text = ""

    if first_bold_idx is not None and first_nonbold_idx is not None:
        para.runs[first_bold_idx].text = new_category + " "
        para.runs[first_nonbold_idx].text = new_skills
    elif first_bold_idx is not None:
        # All runs are bold - put category in first, skills after a space
        # but ensure skills aren't displayed as bold by removing bold from last run
        para.runs[first_bold_idx].text = new_category + " "
        last_idx = len(para.runs) - 1
        if last_idx > first_bold_idx:
            para.runs[last_idx].text = new_skills
            para.runs[last_idx].bold = False
        else:
            para.runs[first_bold_idx].text = new_category + " " + new_skills
    else:
        para.runs[0].text = new_category + " " + new_skills


def update_resume(doc: Document, token: str, jd_text: str) -> Document:
    exp_blocks = build_experience_blocks(doc)
    print(f"Found {len(exp_blocks)} experience blocks")

    for block in exp_blocks:
        company = block["company_text"]
        title = block.get("title_text", "")
        bullet_indices = block["bullet_indices"]

        if not bullet_indices:
            continue

        original_bullets = [doc.paragraphs[idx].text.strip() for idx in bullet_indices]
        print(f"\nTailoring: {title} at {company}")
        print(f"  {len(original_bullets)} bullets to rewrite...")

        new_bullets = tailor_bullets(token, jd_text, company, title, original_bullets)

        for idx, new_text in zip(bullet_indices, new_bullets):
            replace_bullet_text(doc.paragraphs[idx], new_text)
            print(f"  [OK] Updated bullet {idx}")

    # Tailor skills section
    skill_entries = build_skills_entries(doc)
    if skill_entries:
        print(f"\nTailoring skills section ({len(skill_entries)} categories)...")
        for e in skill_entries:
            print(f"  Original: {e['category']} {e['skills'][:60]}...")

        new_entries = tailor_skills_entries(token, jd_text, skill_entries)

        for orig_entry, new_entry in zip(skill_entries, new_entries):
            idx = orig_entry["para_idx"]
            replace_skills_text(
                doc.paragraphs[idx],
                new_entry["category"],
                new_entry["skills"]
            )
            print(f"  [OK] Updated skill para {idx}: {new_entry['category']}")

            # Clear any continuation paragraphs (their content is now
            # merged into the main category line by the LLM)
            for cont_idx in orig_entry.get("continuation_indices", []):
                for run in doc.paragraphs[cont_idx].runs:
                    run.text = ""
                print(f"  [OK] Cleared continuation para {cont_idx}")

    return doc


def convert_to_pdf(docx_path: str, pdf_path: str):
    """Convert DOCX to PDF using LibreOffice (available in GitHub Actions)."""
    import subprocess
    output_dir = str(Path(pdf_path).parent)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", output_dir, docx_path],
        capture_output=True, text=True, timeout=120
    )
    if result.returncode != 0:
        print(f"LibreOffice PDF conversion failed: {result.stderr}")
        soffice_result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf",
             "--outdir", output_dir, docx_path],
            capture_output=True, text=True, timeout=120
        )
        if soffice_result.returncode != 0:
            print(f"soffice also failed: {soffice_result.stderr}")
            raise RuntimeError("PDF conversion failed")

    expected_pdf = Path(output_dir) / (Path(docx_path).stem + ".pdf")
    if expected_pdf.exists() and str(expected_pdf) != pdf_path:
        expected_pdf.rename(pdf_path)


def main():
    parser = argparse.ArgumentParser(description="Tailor resume to a job description")
    parser.add_argument("--jd", required=True, help="Path to the job description file")
    parser.add_argument("--template", default="template/resume_template.docx",
                        help="Path to the resume template DOCX")
    parser.add_argument("--output-dir", default="output", help="Output directory")
    parser.add_argument("--token", default=None,
                        help="GitHub token (or set GITHUB_TOKEN env var)")
    args = parser.parse_args()

    token = args.token or os.environ.get("GH_MODELS_TOKEN") or os.environ.get("GITHUB_TOKEN")
    if not token:
        sys.exit("ERROR: Provide --token or set GH_MODELS_TOKEN environment variable")

    jd_path = Path(args.jd)
    if not jd_path.exists():
        sys.exit(f"ERROR: JD file not found: {jd_path}")

    template_path = Path(args.template)
    if not template_path.exists():
        sys.exit(f"ERROR: Template not found: {template_path}")

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    jd_name = jd_path.stem.replace(" ", "_")
    output_docx = output_dir / f"Rehan_Malik_Resume_{jd_name}.docx"
    output_pdf = output_dir / f"Rehan_Malik_Resume_{jd_name}.pdf"

    print(f"Reading JD: {jd_path}")
    jd_text = read_jd(str(jd_path))
    print(f"JD length: {len(jd_text)} characters")

    print(f"Loading template: {template_path}")
    doc = Document(str(template_path))

    print("Tailoring resume...")
    updated_doc = update_resume(doc, token, jd_text)

    print(f"Saving DOCX: {output_docx}")
    updated_doc.save(str(output_docx))

    print(f"Converting to PDF: {output_pdf}")
    try:
        convert_to_pdf(str(output_docx), str(output_pdf))
        print("PDF created successfully!")
    except Exception as e:
        print(f"PDF conversion skipped ({e}). DOCX is still available.")

    print("\nDone!")
    print(f"  DOCX: {output_docx}")
    print(f"  PDF:  {output_pdf}")


if __name__ == "__main__":
    main()
